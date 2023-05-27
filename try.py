


import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from docx import Document

# Configure Selenium options
chrome_options = Options()
chrome_options.add_argument("--headless")  # Run Chrome in headless mode, without opening a browser window

# Path to your Chrome WebDriver executable
webdriver_path = r"C:\Users\harry\Downloads\chromedriver_win32"

# Create a new Chrome browser instance
browser = webdriver.Chrome(executable_path=webdriver_path, options=chrome_options)

# Navigate to the webpage
url = "https://parade.com/1048764/marynliles/bible-trivia-questions/"
browser.get(url)

# Wait for the page to load (adjust the sleep time if needed)
time.sleep(5)

# Get the page source after it has fully loaded
page_source = browser.page_source

# Create a BeautifulSoup object to parse the HTML content
soup = BeautifulSoup(page_source, "html.parser")

# Find all the text within the <p> tags on the page
paragraphs = soup.find_all("p")

# Create a new Word document
doc = Document()

# Iterate through the paragraphs and add them to the document
for paragraph in paragraphs:
    text = paragraph.get_text(strip=True)
    doc.add_paragraph(text)

# Save the Word document
doc.save("webpage_content.docx")

# Quit the browser
browser.quit()



# # Define the URL of the website you want to scrape
# url = "https://parade.com/1048764/marynliles/bible-trivia-questions/"

# headers = {
#     "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.93 Safari/537.36",
#     "Accept-Language": "en-US,en;q=0.9",
#     "Referer": "https://www.google.com/"
# }

# req = Request( url= url, headers=headers)
# webpage = urlopen(req).read()
# # url_1 =  urllib.request.urlopen(url).read()
# # Send a GET request to the website
# # response = requests.get(url)

# # Create a BeautifulSoup object to parse the HTML content
# soup = BeautifulSoup(webpage, "lxml")
# type(soup)

# # Find the elements containing the text you want to scrape
# text_elements = soup.find_all("p")  # Example: Find all <p> elements

# for t in text_elements:
#     print(t.get_text().strip())
    
# # Create a new document
# doc = Document()

# # Iterate over the text elements and extract the text
# for element in text_elements:
#     text = element.get_text()
#     doc.add_paragraph(text)

# # Save the document as a text file
# doc.save("scraped_text.docx")




# from selenium import webdriver
# from bs4 import BeautifulSoup

# # Specify the path to your webdriver (e.g., ChromeDriver)
# webdriver_path = r"C:\Users\harry\Downloads\chromedriver_win32"

# # Create a new Chrome driver instance
# driver = webdriver.Chrome()

# # Navigate to the website
# url = "https://parade.com/1048764/marynliles/bible-trivia-questions/"
# driver.get(url)

# # Wait for the page to load (you can increase the sleep duration if needed)
# import time
# time.sleep(5)  # Wait for 5 seconds

# # Get the page source after JavaScript execution
# page_source = driver.page_source

# # Create a BeautifulSoup object to parse the HTML content
# soup = BeautifulSoup(page_source, "html.parser")

# # Find the elements containing the text you want to scrape
# text_elements = soup.find_all("p")  # Example: Find all <p> elements

# # Extract and print the text from the elements
# for element in text_elements:
#     text = element.get_text()
#     print(text)

# # Close the browser
# driver.quit()


# p = doc.add_paragraph('This is the start of the paragraph')
# run = p.add_run()
# run.add_break(doc.text.run.WD_BREAK.LINE)
# p.add_run('And now this in a different line')
# p.add_run(". Even if it's on the same paragraph.")