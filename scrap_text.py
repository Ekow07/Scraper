import requests
# from selenium import webdriver
from urllib.request import Request, urlopen
# from IPython.display import HTML
from bs4 import BeautifulSoup
from docx import Document


# Send a GET request to the webpage
url = "https://ministry-to-children.com/kids-bible-trivia/"

# headers = {
#     "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
# }
response = requests.get(url,)
response.raise_for_status()

# Create a BeautifulSoup object to parse the HTML content
soup = BeautifulSoup(response.text, "html.parser")

print(type(soup))
# Find all the text within the <p> tags on the page
paragraphs = soup.find_all("p")

# Create a new Word document
doc = Document()

# # Iterate through the paragraphs and add them to the document
for paragraph in paragraphs:
    text = paragraph.get_text(strip=True)
    doc.add_paragraph(text)
    doc.add_paragraph("")
    
    
    # print(text)

# Save the Word document
doc.save("try.docx")

